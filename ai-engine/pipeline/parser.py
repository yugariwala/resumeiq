"""
Step 1 — Document Parser
Parse PDF (pypdf) and DOCX (python-docx) files to extract raw text.
"""

from __future__ import annotations
import io

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def parse_document(file_bytes: bytes, file_format: str) -> dict:
    """
    Parse a resume file and extract text.

    Args:
        file_bytes: Raw file bytes
        file_format: 'pdf' or 'docx'

    Returns:
        dict with raw_text, page_count, is_image_only, format
        OR error dict with code and message
    """
    if file_format == "pdf":
        return _parse_pdf(file_bytes)
    elif file_format == "docx":
        return _parse_docx(file_bytes)
    elif file_format == "doc":
        return {
            "code": "OLD_FORMAT",
            "message": "The .doc format is not supported. Please convert to .docx first (File → Save As → .docx in Word).",
        }
    else:
        return {
            "code": "UNSUPPORTED_FORMAT",
            "message": f"Unsupported file format: {file_format}. Please upload PDF or DOCX.",
        }


def _parse_pdf(file_bytes: bytes) -> dict:
    """Extract text from PDF using pypdf."""
    if not HAS_PYPDF:
        return {"code": "MISSING_DEP", "message": "pypdf not installed. Run: pip install pypdf"}

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        return {"code": "PARSE_FAILED", "message": f"Could not open PDF: {str(e)}"}

    # W8: Detect password-protected PDFs
    if reader.is_encrypted:
        return {
            "code": "ENCRYPTED",
            "message": "This PDF is password-protected. Please remove the password and re-upload.",
        }

    page_count = len(reader.pages)
    all_text_parts: list[str] = []
    image_only_pages = 0

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            all_text_parts.append(page_text.strip())
        else:
            image_only_pages += 1

    raw_text = "\n\n".join(all_text_parts)
    is_image_only = image_only_pages == page_count

    # Edge case: too little text
    if not raw_text or len(raw_text.strip()) < 50:
        return {
            "code": "PARSE_FAILED",
            "message": "Resume text could not be extracted. Please upload a text-based PDF.",
        }

    return {
        "raw_text": raw_text,
        "page_count": page_count,
        "is_image_only": is_image_only,
        "format": "pdf",
    }


def _parse_docx(file_bytes: bytes) -> dict:
    """Extract text from DOCX using python-docx."""
    if not HAS_DOCX:
        return {"code": "MISSING_DEP", "message": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        return {"code": "PARSE_FAILED", "message": f"Could not open DOCX: {str(e)}"}

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    raw_text = "\n".join(parts)

    if not raw_text or len(raw_text.strip()) < 50:
        return {
            "code": "PARSE_FAILED",
            "message": "Resume text could not be extracted. Please upload a text-based DOCX.",
        }

    # Estimate page count (roughly 500 words per page)
    word_count = len(raw_text.split())
    estimated_pages = max(1, round(word_count / 500))

    return {
        "raw_text": raw_text,
        "page_count": estimated_pages,
        "is_image_only": False,
        "format": "docx",
    }
